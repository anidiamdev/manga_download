from importlib.resources import path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def do_docx_logic(data):

    document = create_docx(data)
    document = add_images(data, document)
    save_docx(data, document)

    return 1


def create_docx(data):
    """"
    - Creates the word document
    - Adds it a header with the manga's name, and the downloaded chapters
    """
    # getting the data
    manga_title = data["manga_title"]
    chapters = data["chapters"]

    # creating the document and adding the header
    document = Document()
    header = document.sections[0].header
    header_text = header.paragraphs[0]

    # creating the text the header is going to use
    text_for_header = "{}\t(chapter {})".format(manga_title, chapters[0])

    # if there is more than one chapter the text must be different
    if len(chapters) > 1:

        text_for_header = "{}\t(chapters {} to {})".format(manga_title, chapters[0], chapters[-1])
    
    # styling the header (center alignment - bold text)
    header_text.add_run(text_for_header).bold = True
    header_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    return document  # will use it for adding the pictures


def add_images(data, document):
    """
    adds the images to the main document
    and saves it to a provisional path
    """
    paths_to_downloaded_images = data["paths_to_downloaded_images"]

    for chapter_number in paths_to_downloaded_images:

        paragraph = document.add_paragraph()
        paragraph.add_run(text="chapter #{}".format(chapter_number)).bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for image_path in paths_to_downloaded_images[chapter_number]:

            document.add_picture(image_path)
            last_paragraph = document.paragraphs[-1] 
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    return document


def save_docx(data, document):

    manga_title = data["manga_title"]
    chapters = data["chapters"]

    path = "{}\\mangas\\{}\\{} ({} - {}).docx".format(os.getcwd(), manga_title, manga_title, chapters[0], chapters[-1])

    document.save(path)


if __name__ == "__main__":

    # just some testing stuff
    doc = do_docx_logic({
        'manga_title': 'shingeki no kyojin before the fall', 
        'chapters': [1],
        'document_type': 'docx'
        })

