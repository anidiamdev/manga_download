from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def do_docx_logic(data):

    document = create_docx(data)
    add_images(data, document)


def create_docx(data):
    """"
    - Creates the word document
    - Adds it a header with the manga's name, and the downloaded chapters
    """

    # getting the data
    manga_title = data["manga_title"]
    chapters = data["chapters"]

    # creating the document and adding the header
    document = Document()
    header = document.sections[0].header
    header_text = header.paragraphs[0]

    # creating the text the header is going to use
    text_for_header = "{}\t(chapter {})".format(manga_title, chapters[0])

    # if there is more than one chapter the text must be different
    if len(chapters) > 1:

        text_for_header = "{}\t(chapters {} to {})".format(manga_title, chapters[0], chapters[-1])
    
    # styling the header (center alignment - bold text)
    header_text.add_run(text_for_header).bold = True
    header_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    return document  # will use it for adding the pictures


def add_images(data, document):
    """
    adds the images to the main document
    """
    paths_to_downloaded_images = data["paths_to_downloaded_images"]

    return 0


if __name__ == "__main__":

    # just some testing stuff
    doc = create_docx({
        'manga_title': 'shingeki no kyojin before the fall', 
        'chapters': [1, 2, 3]
        })
    
    # saves the document for testing
    doc.save(os.getcwd() + "\\test_document.docx")
